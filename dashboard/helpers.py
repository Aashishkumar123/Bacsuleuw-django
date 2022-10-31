from docx import Document as Doc
import os
import aspose.words as aw


def generate_document(agency_name,agency_address,agency_state,agency_city,agency_zip,date,agent_name,agency_commission,email):
    filename = os.path.join('static','documents/Blank-Bascule-Agency-Agreement-Equine.docx')
    doc = Doc(filename)
    for p in doc.paragraphs:
       inline = p.runs
       for i in range(len(inline)):
            if inline[i].text == '(Agency Name)' or inline[i].text == ' (Agency Name) ':
                inline[i].text = agency_name
            if inline[i].text == '(Agent Address)':
                inline[i].text = f'{agency_address} {agency_state}, {agency_city}, {agency_zip}'
            if inline[i].text == '(Agent Street Address)':
                inline[i].text = agency_address
            if inline[i].text == '(Agent City, State, Zip)':
                inline[i].text = f'{agency_state}, {agency_city}, {agency_zip}'
            if inline[i].text == '(Date)':
                inline[i].text = '20-09-1999'
            if inline[i].text == '(Agent Name)':
                inline[i].text = {agent_name}
            if inline[i].text == '(Agent Address),':
                inline[i].text = f'{agency_address} {agency_state}, {agency_city}, {agency_zip}'
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '__20_%' in paragraph.text:
                        paragraph.text = f"As negotiated but not to exceed __{agency_commission}_%"
    filename = f'static/documents/Blank-Bascule-Agency-Agreement-Equine-{email}'
    doc.save(f'{filename}.docx')
    doc = aw.Document(f'{filename}.docx')
    doc.save(f'{filename}.pdf')
